# import PyPDF2  # 用于读取PDF文件
import docx    # 用于读取Word文档
import os      # 用于文件路径操作
import re      # 用于正则表达式
from abc import ABC, abstractmethod  # 用于创建抽象基类

# 定义文件读取器的抽象基类
class FileReader(ABC):
    @abstractmethod
    def read(self, file_path):
        """读取文件的抽象方法"""
        pass

# PDF文件读取器
class PDFReader(FileReader):
    def read(self, file_path):
        """读取PDF文件并提取文本内容"""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            return '\n'.join(page.extract_text() for page in pdf_reader.pages)

# TXT文本文件读取器
class TXTReader(FileReader):
    def read(self, file_path):
        """读取文本文件内容"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

# # Word文档读取器
# class WordReader(FileReader):
#     def read(self, file_path):
#         """读取Word文档并提取文本内容"""
#         doc = docx.Document(file_path)
#         return '\n'.join(para.text for para in doc.paragraphs)

# Word 文档读取器，返回树状结构
class WordReader(FileReader):
    def read(self, file_path):
        """读取 Word 文档并返回树状结构：标题+段落结构"""
        doc = docx.Document(file_path)
        root = []
        stack = []

        def get_heading_level(style_name):
            match = re.match(r"Heading (\d+)", style_name)
            return int(match.group(1)) if match else None

        def new_section(title):
            return {"title": title, "paragraphs": [], "children": []}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            level = get_heading_level(para.style.name)
            if level:  # 是标题
                section = new_section(text)
                while stack and stack[-1][0] >= level:
                    stack.pop()
                if stack:
                    stack[-1][1]["children"].append(section)
                else:
                    root.append(section)
                stack.append((level, section))
            else:
                if stack:
                    stack[-1][1]["paragraphs"].append(text)
                else:
                    if not root or "title" in root[-1]:
                        root.append(new_section("无标题内容"))
                    root[-1]["paragraphs"].append(text)
        return root

# 文件读取器工厂类
class FileReaderFactory:
    # 支持的文件类型及对应的读取器实例
    readers = {
        '.pdf': PDFReader(),
        '.txt': TXTReader(),
        '.doc': WordReader(),
        '.docx': WordReader()
    }

    @staticmethod
    def get_reader(file_ext):
        """根据文件扩展名获取对应的读取器"""
        reader = FileReaderFactory.readers.get(file_ext.lower())
        if not reader:
            raise ValueError(f'不支持的文件格式: {file_ext}')
        return reader
    
# 文件读取类
class ReadFile:
    def __init__(self, file_path):
        """初始化文件路径和扩展名"""
        self.file_path = file_path
        self.file_ext = os.path.splitext(file_path)[-1].lower()

    def get_file_content(self):
        """获取文件内容"""
        reader = FileReaderFactory.get_reader(self.file_ext)
        return reader.read(self.file_path)


if __name__ == '__main__':
    # 测试文件读取
    file_path = '/home/baishi/PycharmProjects/text_gen/test/testdoc.docx'
    reader = ReadFile(file_path)
    try:
        content = reader.get_file_content()
        for i in content:

        
            print(f'文件内容:\n{i}')
    except Exception as e:
        print(f'读取文件时发生错误: {str(e)}')
